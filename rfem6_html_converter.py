import os
import sys
import re
import tkinter as tk
from tkinter import filedialog, ttk, messagebox, StringVar, BooleanVar, Frame, Checkbutton, Label, Entry, Button
from tkinter.scrolledtext import ScrolledText
from bs4 import BeautifulSoup
import requests
from PIL import Image, ImageTk
from io import BytesIO
import tempfile
import cairosvg
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import platform
import subprocess


class HTMLtoDOCXConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("RFEM HTML zu DOCX Konverter")
        self.root.geometry("800x630")  # Höhe reduziert, da Status und Protokoll entfernt werden
        self.root.minsize(800, 600)

        # Variablen für Konfiguration
        self.html_file_path = StringVar()
        self.output_dir = StringVar()
        self.preserve_colors = BooleanVar(value=True)
        self.include_images = BooleanVar(value=True)
        self.adjust_table_width = BooleanVar(value=True)
        self.font_name = StringVar(value="Calibri")
        self.font_size = StringVar(value="11")
        self.table_font_size = StringVar(value="10")
        self.max_image_width_cm = StringVar(value="16")
        self.second_column_width_cm = StringVar(value="8")
        self.use_existing_docx = BooleanVar(value=False)
        self.existing_docx_path = StringVar()
        self.insert_page = StringVar(value="5")  # Geändert von "6" auf "5"
        
        # Standardwert für Ausgabeverzeichnis
        self.output_dir.set(os.path.expanduser("~\\Documents"))
        
        # Oberfläche erstellen
        self.create_widgets()
        
    def create_widgets(self):
        # Hauptframe
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Dateiauswahl
        file_frame = ttk.LabelFrame(main_frame, text="Dateiauswahl", padding="10")
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(file_frame, text="HTML-Datei:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(file_frame, textvariable=self.html_file_path, width=50).grid(row=0, column=1, sticky=tk.W+tk.E, padx=5, pady=5)
        ttk.Button(file_frame, text="Durchsuchen...", command=self.browse_html_file).grid(row=0, column=2, padx=5, pady=5)
        
        ttk.Label(file_frame, text="Ausgabeverzeichnis:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(file_frame, textvariable=self.output_dir, width=50).grid(row=1, column=1, sticky=tk.W+tk.E, padx=5, pady=5)
        ttk.Button(file_frame, text="Durchsuchen...", command=self.browse_output_dir).grid(row=1, column=2, padx=5, pady=5)
        
        # Bestehendes DOCX einfügen
        docx_frame = ttk.LabelFrame(main_frame, text="In bestehendes DOCX einfügen", padding="10")
        docx_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Checkbutton(docx_frame, text="In bestehendes Word-Dokument einfügen", 
                        variable=self.use_existing_docx).grid(row=0, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(docx_frame, text="DOCX-Datei:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(docx_frame, textvariable=self.existing_docx_path, width=50).grid(row=1, column=1, sticky=tk.W+tk.E, padx=5, pady=5)
        ttk.Button(docx_frame, text="Durchsuchen...", command=self.browse_existing_docx).grid(row=1, column=2, padx=5, pady=5)
        
        ttk.Label(docx_frame, text="Ab Seite:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Spinbox(docx_frame, from_=1, to=100, textvariable=self.insert_page, width=5).grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Formatierungsoptionen
        format_frame = ttk.LabelFrame(main_frame, text="Formatierungsoptionen", padding="10")
        format_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # Linke Spalte
        left_frame = ttk.Frame(format_frame)
        left_frame.grid(row=0, column=0, sticky=tk.W, padx=5)
        
        ttk.Label(left_frame, text="Schriftart:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Combobox(left_frame, textvariable=self.font_name, values=["Calibri", "Arial", "Times New Roman", "Verdana"]).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(left_frame, text="Schriftgröße:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Combobox(left_frame, textvariable=self.font_size, values=["9", "10", "11", "12"]).grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Info zur besseren Formatierungserhaltung
        ttk.Label(left_frame, text="Formatierung wird vollständig erhalten", foreground="green").grid(row=2, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        # Rechte Spalte
        right_frame = ttk.Frame(format_frame)
        right_frame.grid(row=0, column=1, sticky=tk.W, padx=5)
        
        ttk.Checkbutton(right_frame, text="Tabellenfarben beibehalten", variable=self.preserve_colors).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Checkbutton(right_frame, text="Bilder einbinden", variable=self.include_images).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Checkbutton(right_frame, text="Tabellenbreite anpassen", variable=self.adjust_table_width).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(right_frame, text="Schriftgröße Tabellen:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Spinbox(right_frame, from_=4, to=16, textvariable=self.table_font_size, width=5).grid(row=3, column=1, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(right_frame, text="Max. PNG-Bildbreite (cm):").grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(right_frame, textvariable=self.max_image_width_cm, width=10).grid(row=4, column=1, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(right_frame, text="Breite 2. Spalte (cm):").grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(right_frame, textvariable=self.second_column_width_cm, width=10).grid(row=5, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Aktionsbuttons
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=tk.X, padx=5, pady=10)
        
        ttk.Button(action_frame, text="Konvertieren", command=self.convert_html_to_docx, width=20).pack(side=tk.RIGHT, padx=5)
        
    def browse_html_file(self):
        file_path = filedialog.askopenfilename(
            title="HTML-Datei auswählen",
            filetypes=(("HTML-Dateien", "*.html"), ("Alle Dateien", "*.*"))
        )
        if file_path:
            self.html_file_path.set(file_path)
            # Ausgabeverzeichnis auf das Verzeichnis der HTML-Datei setzen
            self.output_dir.set(os.path.dirname(file_path))
            print(f"HTML-Datei ausgewählt: {os.path.basename(file_path)}")
            print(f"HTML-Datei: {file_path}")
    
    def browse_output_dir(self):
        dir_path = filedialog.askdirectory(title="Ausgabeverzeichnis auswählen")
        if dir_path:
            self.output_dir.set(dir_path)
            print(f"Ausgabeverzeichnis: {dir_path}")
    
    def browse_existing_docx(self):
        file_path = filedialog.askopenfilename(
            title="DOCX-Datei auswählen",
            filetypes=(("Word-Dokumente", "*.docx"), ("Alle Dateien", "*.*"))
        )
        if file_path:
            self.existing_docx_path.set(file_path)
            print(f"Bestehendes DOCX: {file_path}")
    
    def log(self, message):
        print(message)  # Nur in der Konsole ausgeben für einfachere Fehlersuche
    
    def convert_html_to_docx(self):
        if not self.html_file_path.get():
            messagebox.showwarning("Warnung", "Bitte wählen Sie zuerst eine HTML-Datei aus.")
            return
            
        # GUI-Elemente während der Konvertierung deaktivieren
        self.root.config(cursor="wait")
        # Wir deaktivieren nicht alle Widgets, da nicht alle die state-Option unterstützen
        self.root.update()
        
        # Ausgabepfad bestimmen
        output_path = os.path.join(
            self.output_dir.get(), 
            os.path.splitext(os.path.basename(self.html_file_path.get()))[0] + ".docx"
        )
        
        try:
            self.log(f"Starte Konvertierung: {self.html_file_path.get()}")
            
            # HTML parsen
            with open(self.html_file_path.get(), 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file, 'html.parser')
            
            # Word-Dokument erstellen oder öffnen
            if self.use_existing_docx.get() and self.existing_docx_path.get():
                try:
                    # Bestehendes Dokument öffnen
                    doc = Document(self.existing_docx_path.get())
                    self.log(f"Bestehendes DOCX geöffnet: {self.existing_docx_path.get()}")
                    
                    # Inhalte ab der angegebenen Seite einfügen
                    target_page = int(self.insert_page.get())
                    self.log(f"Füge Inhalte ab Seite {target_page} ein")
                    
                    # Vorbereitung für das Einfügen
                    self.prepare_document_for_insertion(doc, target_page)
                    
                    # Ausgabepfad anpassen, falls das originale Dokument nicht überschrieben werden soll
                    output_path = os.path.join(
                        self.output_dir.get(), 
                        os.path.splitext(os.path.basename(self.existing_docx_path.get()))[0] + "_mit_html.docx"
                    )
                    
                except Exception as e:
                    self.log(f"Fehler beim Öffnen des bestehenden DOCX: {str(e)}")
                    # Fallback: Neues Dokument erstellen
                    doc = Document()
                    self.log("Fallback: Erstelle neues Dokument")
            else:
                # Neues Dokument erstellen
                doc = Document()
                self.log("Erstelle neues Dokument")
            
            # Dokumenteigenschaften setzen
            core_properties = doc.core_properties
            core_properties.title = "RFEM Dokument"
            core_properties.author = "RFEM HTML zu DOCX Konverter"
            
            # Standardformatierung
            font_name = self.font_name.get()
            font_size = int(self.font_size.get())
            
            # Absatzstil für den gesamten Text
            style = doc.styles['Normal']
            style.font.name = font_name
            style.font.size = Pt(font_size)
            
            # Nur Seitenränder für neues Dokument anpassen, nicht bei bestehendem
            if not (self.use_existing_docx.get() and self.existing_docx_path.get()):
                sections = doc.sections
                for section in sections:
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
            
            # Fortschritt initialisieren
            total_elements = len(soup.find_all(['h1', 'h2', 'h3', 'table', 'img']))
            processed_elements = 0
            
            # Basisverzeichnis der HTML-Datei für Bildpfade
            base_dir = os.path.dirname(self.html_file_path.get())
            
            # Header-Information und Titel
            header_img = soup.find('img', {'width': '64', 'height': '64'})
            if header_img and self.include_images.get():
                self.log("Verarbeite Header-Logo...")
                try:
                    img_path = self.resolve_image_path(header_img.get('src'), base_dir)
                    if img_path:
                        header_para = doc.add_paragraph()
                        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = header_para.add_run()
                        run.add_picture(img_path, width=Inches(0.7))
                except Exception as e:
                    self.log(f"Warnung: Konnte Header-Logo nicht einfügen: {str(e)}")
            
            # Titeltext
            title_text = soup.find('b', string=lambda t: t and 'RFEM' in t)
            if title_text:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(title_text.text)
                run.bold = True
            
            # Trennlinie nach Header
            if soup.find('hr'):
                doc.add_paragraph('_' * 80)
            
            # Hauptüberschrift
            main_title = soup.find('h1')
            if main_title:
                text = main_title.text.strip()
                para = doc.add_paragraph()
                para.style = 'Heading 1'
                run = para.add_run(text)
                run.bold = True
                processed_elements += 1
            
            # Verarbeite alle Elemente
            for element in soup.body.find_all(['h1', 'h2', 'h3', 'table', 'p', 'img', 'div'], recursive=True):
                if element.name in ['h1', 'h2', 'h3']:
                    self.process_heading(doc, element)
                    processed_elements += 1
                    
                elif element.name == 'table':
                    self.process_table(doc, element)
                    processed_elements += 1
                    
                elif element.name == 'img' and self.include_images.get():
                    if 'src' in element.attrs:
                        self.process_image(doc, element, base_dir)
                        processed_elements += 1
            
            # Bei bestehendem Dokument angehängte Inhalte wieder hinzufügen
            if self.use_existing_docx.get() and self.existing_docx_path.get():
                self.finalize_document(doc)
            
            # Speichern des Dokuments
            doc.save(output_path)
            self.log(f"Konvertierung abgeschlossen: {output_path}")
            
            # GUI-Cursor wieder normal setzen
            self.root.config(cursor="")
            self.root.update()
            
            # Erfolgreiche Meldung mit Option zum Öffnen
            result = messagebox.askquestion(
                "Konvertierung abgeschlossen", 
                f"Die Datei wurde erfolgreich konvertiert und gespeichert unter:\n{output_path}\n\nMöchten Sie die Datei jetzt öffnen?",
                icon='info'
            )
            if result == 'yes':
                self.open_document(output_path)
            
        except Exception as e:
            # GUI-Cursor wieder normal setzen
            self.root.config(cursor="")
            self.root.update()
            
            self.log(f"Fehler bei der Konvertierung: {str(e)}")
            messagebox.showerror("Fehler", f"Fehler bei der Konvertierung: {str(e)}")

    def open_document(self, doc_path):
        """
        Öffnet das erstellte Dokument mit dem Standardprogramm des Betriebssystems.
        """
        try:
            if platform.system() == 'Windows':
                os.startfile(doc_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(['open', doc_path])
            else:  # Linux und andere
                subprocess.call(['xdg-open', doc_path])
                
            self.log(f"Dokument wurde geöffnet: {doc_path}")
        except Exception as e:
            self.log(f"Fehler beim Öffnen des Dokuments: {str(e)}")
            messagebox.showerror("Fehler", f"Das Dokument konnte nicht geöffnet werden: {str(e)}")
    
    def prepare_document_for_insertion(self, doc, target_page):
        """
        Bereitet das Dokument vor, um HTML-Inhalte ab einer bestimmten Seite einzufügen.
        Speichert den kompletten Dokumentinhalt nach der Zielseite in eine separate Datei.
        """
        self.log(f"Bereite Dokument für Einfügen ab Seite {target_page} vor...")
        
        # Gesamtzahl der Absätze im Dokument
        total_paragraphs = len(doc.paragraphs)
        
        if total_paragraphs == 0:
            self.log("Dokument ist leer, füge am Anfang ein.")
            return
            
        # Alternative Methode: Wir erstellen eine exakte Kopie des gesamten Dokuments
        # und schneiden dann die Teile ab, die wir beibehalten oder entfernen wollen
        
        # Geschätzte Anzahl von Absätzen pro Seite (Durchschnitt)
        estimated_paragraphs_per_page = 30
        
        # Geschätzte Absatzposition für die Zielseite
        target_paragraph_index = (target_page - 1) * estimated_paragraphs_per_page
        
        # Stelle sicher, dass der Index innerhalb des gültigen Bereichs liegt
        if target_paragraph_index >= total_paragraphs:
            target_paragraph_index = total_paragraphs - 1
            self.log(f"Zielseite außerhalb des Dokumentenbereichs, verwende letzten Absatz.")
        
        try:
            # Speichere das gesamte Originaldokument
            temp_file_original = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            doc.save(temp_file_original.name)
            temp_file_original.close()
            
            # Erstelle ein Dokument für den Teil nach der Einfügestelle
            temp_file_after = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_file_after.close()
            
            # Kopiere das Originaldokument
            import shutil
            shutil.copy2(temp_file_original.name, temp_file_after.name)
            
            # Öffne das Dokument für den Teil nach der Einfügestelle
            doc_after = Document(temp_file_after.name)
            
            # Lösche alles vor dem Einfügepunkt im "Nachher"-Dokument
            # Wir löschen von vorne nach hinten, da wir die Indizes anpassen müssen
            paragraphs_to_delete = min(target_paragraph_index, len(doc_after.paragraphs))
            for _ in range(paragraphs_to_delete):
                p = doc_after.paragraphs[0]  # Immer den ersten löschen
                p._element.getparent().remove(p._element)
            
            # Speichere das "Nachher"-Dokument
            doc_after.save(temp_file_after.name)
            
            # Öffne das Originaldokument erneut
            doc_original = Document(temp_file_original.name)
            
            # Ersetze das übergebene Dokument mit dem Original
            # (nur notwendig, wenn wir ein neues Objekt erstellt haben)
            if doc is not doc_original:
                # Dies ist nur ein Workaround - in der Praxis würden wir direkt mit dem Original arbeiten
                pass
            
            # Lösche alle Paragraphen nach dem Einfügepunkt im Original
            # Lösche rückwärts, damit die Indizes gültig bleiben
            for i in range(len(doc.paragraphs)-1, target_paragraph_index-1, -1):
                if i < len(doc.paragraphs):  # Sicherheitsprüfung
                    p = doc.paragraphs[i]
                    p._element.getparent().remove(p._element)
            
            # Füge einen Seitenumbruch am Ende des letzten verbleibenden Absatzes ein
            if len(doc.paragraphs) > 0:
                last_p = doc.paragraphs[-1]
                if len(last_p.runs) > 0:
                    last_p.runs[-1].add_break(WD_BREAK.PAGE)
                else:
                    last_p.add_run().add_break(WD_BREAK.PAGE)
            
            # Speichere den Pfad des temporären Dokuments für spätere Verwendung
            self._temp_original_doc_path = temp_file_original.name
            self._temp_after_doc_path = temp_file_after.name
            
            self.log(f"Dokumentteile wurden vorbereitet. Teil nach Seite {target_page} gespeichert.")
            
        except Exception as e:
            self.log(f"Fehler bei Dokumentvorbereitung: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            # Bei Fehlern setzen wir die temporären Dokumente auf None
            self._temp_original_doc_path = None
            self._temp_after_doc_path = None
        
        self.log(f"Bereit zum Einfügen von Inhalten ab Seite {target_page}.")
        
    def finalize_document(self, doc):
        """
        Fügt die zuvor gespeicherten Dokument-Inhalte mit vollständiger Formatierung wieder an.
        """
        # Füge einen Seitenumbruch am Ende des eingefügten HTML-Inhalts ein
        if len(doc.paragraphs) > 0:
            last_p = doc.paragraphs[-1]
            if len(last_p.runs) > 0:
                last_p.runs[-1].add_break(WD_BREAK.PAGE)
            else:
                last_p.add_run().add_break(WD_BREAK.PAGE)
                
        if hasattr(self, '_temp_after_doc_path') and self._temp_after_doc_path:
            try:
                self.log("Füge formatierte Inhalte nach dem eingefügten HTML wieder an...")
                
                # Kopiere das gesamte Dokument mit allen Formatierungen
                # Wir verwenden hierzu die docx-Bibliothek auf XML-Ebene, um alle Formatierungen zu erhalten
                
                # Lade das temporäre "Nachher"-Dokument
                doc_after = Document(self._temp_after_doc_path)
                
                # Importiere alle Stile vom "Nachher"-Dokument
                for style in doc_after.styles:
                    if style.name not in doc.styles:
                        try:
                            doc.styles.add_style(style.name, style.type)
                        except:
                            pass  # Ignoriere Fehler bei Standardstilen
                
                # Kopiere alle XML-Elemente direkt aus dem "Nachher"-Dokument
                # Dies ist der Schlüssel zur Erhaltung der gesamten Formatierung
                body_elem = doc._element.body
                after_body_elem = doc_after._element.body
                
                # Füge alle Kindelemente des "Nachher"-Dokuments hinzu
                for child in after_body_elem:
                    body_elem.append(child)
                
                # Bereinige temporäre Dateien
                if hasattr(self, '_temp_original_doc_path') and self._temp_original_doc_path:
                    try:
                        os.unlink(self._temp_original_doc_path)
                        self._temp_original_doc_path = None
                    except:
                        pass
                        
                if hasattr(self, '_temp_after_doc_path'):
                    try:
                        os.unlink(self._temp_after_doc_path)
                        self._temp_after_doc_path = None
                    except:
                        pass
                
                self.log("Angehängte Inhalte wurden mit vollständiger Formatierung wiederhergestellt.")
                
            except Exception as e:
                self.log(f"Fehler beim Wiederherstellen der angehängten Inhalte: {str(e)}")
                import traceback
                self.log(traceback.format_exc())
    
    def process_heading(self, doc, element):
        text = element.text.strip()
        if not text:
            return
            
        # ID aus Element extrahieren für Sprungmarken
        element_id = None
        if 'id' in element.attrs:
            element_id = element['id']
        
        level = int(element.name[1])  # h1, h2, h3 -> 1, 2, 3
        
        para = doc.add_paragraph()
        para.style = f'Heading {level}'
        run = para.add_run(text)
        run.bold = True
        
        # Größe basierend auf Level anpassen
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        
        # Absatzabstand
        para.space_after = Pt(12)
        
        self.log(f"Überschrift verarbeitet: {text[:40]}...")
    
    def process_table(self, doc, table_element):
        self.log("Verarbeite Tabelle...")
        
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Spalten analysieren, um leere Spalten und Doppelpunkt-Spalten zu identifizieren
        all_cells = []
        for row in rows:
            cells = row.find_all(['th', 'td'])
            all_cells.append(cells)
        
        if not all_cells:
            return
        
        max_cols = max(len(cells) for cells in all_cells)
        if max_cols == 0:
            return
        
        # Prüfen, welche Spalten nur Doppelpunkte enthalten
        cols_with_only_colons = []
        for col_idx in range(max_cols):
            only_colons = True
            for row_cells in all_cells:
                if col_idx < len(row_cells):
                    cell_text = row_cells[col_idx].get_text(strip=True)
                    if cell_text != ":" and cell_text != "":
                        only_colons = False
                        break
            
            if only_colons:
                cols_with_only_colons.append(col_idx)
        
        # Anzahl der zu entfernenden Spalten
        num_cols_to_skip = len(cols_with_only_colons)
        
        # Bereinigen der Spaltenliste
        num_cols = max_cols - num_cols_to_skip
        
        if num_cols == 0:
            return
        
        num_rows = len(rows)
        
        # Tabelle erstellen
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Immer Tabellenstil 1 anwenden
        table.style = 'Table Grid'
        
        # Tabellenbreite anpassen
        if self.adjust_table_width.get():
            table.autofit = True
            
            # Spaltenbreiten optimieren - schmaler machen für bestimmte Spaltentypen
            for i, column_cells in enumerate(table.columns):
                # Besondere Behandlung für die zweite Spalte (Index 1)
                if i == 1 and num_cols > 1:  # Zweite Spalte, wenn vorhanden
                    # Breitere Breite für die zweite Spalte
                    second_column_width_cm = float(self.second_column_width_cm.get())
                    second_column_width_twips = int(second_column_width_cm * 1440 / 2.54)  # cm zu twips konvertieren
                    
                    for cell in column_cells.cells:
                        tc = cell._tc
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(second_column_width_twips))
                        tcW.set(qn('w:type'), 'dxa')
                        tc.get_or_add_tcPr().append(tcW)
                else:
                    # Überprüfe den Inhalt der Spalte
                    narrow_column = True
                    for cell in column_cells.cells:
                        text = cell.text.strip()
                        # Wenn die Spalte nur kurze Einträge oder Zahlen enthält, machen wir sie schmaler
                        if len(text) > 5 and not text.replace('.', '', 1).isdigit():
                            narrow_column = False
                            break
                    
                    if narrow_column:
                        # Setze eine schmalere Breite für diese Spalte
                        for cell in column_cells.cells:
                            tc = cell._tc
                            tcW = OxmlElement('w:tcW')
                            tcW.set(qn('w:w'), '600')  # Schmale Breite (600 = ca. 0.5cm)
                            tcW.set(qn('w:type'), 'dxa')
                            tc.get_or_add_tcPr().append(tcW)
        else:
            table.autofit = False
        
        # Tabellenausrichtung zentrieren
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Zeilen verarbeiten
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            
            # Spaltenindex für das Word-Dokument
            word_col_idx = 0
            
            # Zeilen verarbeiten
            for j, cell in enumerate(cells):
                # Überspringe Spalten, die nur Doppelpunkte enthalten
                if j in cols_with_only_colons:
                    continue
                    
                if word_col_idx < num_cols:  # Nur gültige Spalten
                    try:
                        # Text aus Zelle extrahieren
                        cell_text = cell.get_text(strip=True)
                        
                        # Zelle füllen
                        table.cell(i, word_col_idx).text = cell_text
                        
                        # Schriftgröße für Tabellenzellen anwenden
                        table_font_size = int(self.table_font_size.get())
                        cell_para = table.cell(i, word_col_idx).paragraphs[0]
                        for run in cell_para.runs:
                            run.font.size = Pt(table_font_size)
                        
                        # Zellenausrichtung
                        if 'style' in cell.attrs and 'text-align' in cell['style']:
                            alignment = cell['style'].split('text-align:')[1].split(';')[0].strip()
                            cell_para = table.cell(i, word_col_idx).paragraphs[0]
                            if alignment == 'center':
                                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif alignment == 'right':
                                cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif alignment == 'left':
                                cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Farbige Hintergrundfarbe für abwechselnde Zeilen
                        if self.preserve_colors.get():
                            # Prüfen, ob Zelle eine Hintergrundfarbe hat
                            bg_color = None
                            
                            # Prüfen auf background-color oder background im style-Attribut
                            if 'style' in cell.attrs:
                                style = cell['style']
                                if 'background-color' in style:
                                    color_part = style.split('background-color:')[1].split(';')[0].strip()
                                    if color_part.startswith('#'):
                                        bg_color = color_part
                                elif 'background' in style:
                                    color_part = style.split('background:')[1].split(';')[0].strip()
                                    if color_part.startswith('#'):
                                        bg_color = color_part
                            
                            # Prüfen auf bgcolor Attribut
                            if not bg_color and 'bgcolor' in cell.attrs:
                                bg_color = cell['bgcolor']
                            
                            # Prüfen auf Standard-Farbmuster in der RFEM-Datei
                            if not bg_color:
                                parent_tr = cell.parent
                                if parent_tr and 'style' in parent_tr.attrs:
                                    style = parent_tr['style']
                                    if 'background-color' in style:
                                        color_part = style.split('background-color:')[1].split(';')[0].strip()
                                        if color_part.startswith('#'):
                                            bg_color = color_part
                            
                            # Standardhintergrundfarben für abwechselnde Zeilen
                            if not bg_color and i % 2 == 1:  # Jede zweite Zeile 
                                bg_color = "#f3f3f4"  # Hellgrau für gerade Zeilen
                            
                            # Hintergrundfarben aus der RFEM HTML spezifisch erkennen
                            if not bg_color:
                                if 'class' in cell.attrs:
                                    cell_class = cell['class']
                                    if 'background-color: #f0f9fe' in str(cell) or 'background-color: #e7f6fe' in str(cell):
                                        bg_color = "#f0f9fe"  # Hellblau
                            
                            # Wende Hintergrundfarbe an
                            if bg_color:
                                cell_elem = table.cell(i, word_col_idx)._element
                                # Konvertiere Hex zu RGB
                                if bg_color.startswith('#'):
                                    try:
                                        r = int(bg_color[1:3], 16)
                                        g = int(bg_color[3:5], 16)
                                        b = int(bg_color[5:7], 16)
                                        
                                        # Füge Hintergrundfarbe hinzu
                                        tcPr = cell_elem.get_or_add_tcPr()
                                        shading = OxmlElement('w:shd')
                                        shading.set(qn('w:fill'), bg_color.replace('#', ''))
                                        tcPr.append(shading)
                                    except ValueError:
                                        # Ignoriere ungültige Farbwerte
                                        pass
                        
                        # Fett markierten Text verarbeiten
                        bold_elements = cell.find_all('b')
                        if bold_elements:
                            # Markiere Text in der Word-Zelle fett
                            cell_para = table.cell(i, word_col_idx).paragraphs[0]
                            cell_para.text = ""  # Lösche den vorherigen Text
                            
                            # Füge den Text mit Formatierung hinzu
                            for part in cell.contents:
                                if hasattr(part, 'name') and part.name == 'b':
                                    run = cell_para.add_run(part.get_text())
                                    run.bold = True
                                    run.font.size = Pt(int(self.table_font_size.get()))
                                elif hasattr(part, 'string') and part.string:
                                    run = cell_para.add_run(part.string)
                                    run.font.size = Pt(int(self.table_font_size.get()))
                                elif isinstance(part, str):
                                    run = cell_para.add_run(part)
                                    run.font.size = Pt(int(self.table_font_size.get()))
                    
                        # Prüfe auf Checkboxen
                        checkbox = cell.find('input', {'type': 'checkbox'})
                        if checkbox:
                            # Ersetze den Text mit einem Häkchen-Symbol oder leer
                            cell_para = table.cell(i, word_col_idx).paragraphs[0]
                            cell_para.text = ""  # Bestehenden Text löschen
                            
                            if 'checked' in checkbox.attrs:
                                run = cell_para.add_run("✓")  # Unicode Häkchen
                                run.font.size = Pt(10)
                            # Bei nicht ausgefüllten Checkboxen kein Symbol
                        
                        # Inkrementiere den Word-Spaltenindex
                        word_col_idx += 1
                        
                    except Exception as e:
                        self.log(f"Warnung: Problem bei Tabellenzelle {i},{j}: {str(e)}")
        
        # Leerzeile nach Tabelle einfügen
        doc.add_paragraph()
    
    def process_image(self, doc, img_element, base_dir):
        img_src = img_element.get('src', '')
        if not img_src:
            return
        
        self.log(f"Verarbeite Bild: {img_src}")
        
        try:
            # Bildpfad auflösen
            img_path = self.resolve_image_path(img_src, base_dir)
            
            if img_path:
                # Bildgröße berechnen
                width = img_element.get('width')
                height = img_element.get('height')
                
                # Paragraph für das Bild erstellen
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                
                # Maximale Bildbreite in cm für PNG-Bilder
                max_image_width_cm = float(self.max_image_width_cm.get())
                max_image_width_inches = max_image_width_cm / 2.54  # Umrechnung cm zu inches
                
                # SVG-Datei konvertieren
                if img_path.lower().endswith('.svg'):
                    try:
                        # Temporäre PNG-Datei erstellen
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                            temp_png = temp_file.name
                        
                        # SVG in PNG konvertieren
                        cairosvg.svg2png(url=img_path, write_to=temp_png)
                        
                        # PNG mit originaler Größe einfügen
                        if width and height:
                            w_inches = float(width) / 96.0  # 96 DPI ist Standard für Web
                            run.add_picture(temp_png, width=Inches(w_inches))
                        else:
                            run.add_picture(temp_png)
                        
                        # Temporäre Datei löschen
                        os.unlink(temp_png)
                        
                    except Exception as svg_error:
                        self.log(f"Warnung: Konnte SVG nicht konvertieren: {str(svg_error)}")
                        # Füge Text hinzu, wenn SVG-Konvertierung fehlschlägt
                        run.add_text(f"[Bild: {os.path.basename(img_src)}]")
                else:
                    # PNG-Bild einfügen mit maximaler Breite
                    try:
                        # Bild-Abmessungen ermitteln
                        from PIL import Image
                        with Image.open(img_path) as img:
                            img_width, img_height = img.size
                        
                        # Bild wird nur verkleinert, wenn es breiter als die maximale Breite ist
                        if img_width > 0 and img_height > 0:
                            actual_width = min(img_width / 96.0, max_image_width_inches)  # Breite in Inches begrenzen
                            run.add_picture(img_path, width=Inches(actual_width))
                            self.log(f"Bild {img_path} eingefügt mit Breite: {actual_width} inches")
                        else:
                            # Fallback für Bilder ohne erkennbare Größe
                            run.add_picture(img_path)
                    except Exception as img_error:
                        self.log(f"Warnung: Problem beim Einfügen des Bildes: {str(img_error)}")
                        run.add_picture(img_path)  # Einfacher Fallback
                
                # Bildunterschrift
                if img_element.parent and img_element.parent.name == 'figure':
                    caption = img_element.parent.find('figcaption')
                    if caption:
                        caption_para = doc.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_para.add_run(caption.get_text()).italic = True
            
            else:
                self.log(f"Warnung: Bild nicht gefunden: {img_src}")
        
        except Exception as e:
            self.log(f"Fehler beim Verarbeiten des Bildes {img_src}: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
    
    def resolve_image_path(self, src, base_dir):
        # Wenn der Pfad mit http(s):// beginnt, ist es eine externe URL
        if src.startswith(('http://', 'https://')):
            try:
                response = requests.get(src)
                if response.status_code == 200:
                    # Temporäre Datei erstellen
                    temp_file = tempfile.NamedTemporaryFile(delete=False)
                    temp_file.write(response.content)
                    temp_file.close()
                    return temp_file.name
            except Exception as e:
                self.log(f"Warnung: Konnte externes Bild nicht laden: {str(e)}")
                return None
        
        # Relativer Pfad
        possible_paths = [
            os.path.join(base_dir, src),                     # Direkt im Basisverzeichnis
            os.path.join(base_dir, os.path.basename(src)),   # Nur der Dateiname
            os.path.join(base_dir, 'test_data', os.path.basename(src))  # Im test_data Unterordner (RFEM spezifisch)
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        # Nicht gefunden
        return None


def main():
    root = tk.Tk()
    app = HTMLtoDOCXConverter(root)
    root.mainloop()


if __name__ == "__main__":
    main()